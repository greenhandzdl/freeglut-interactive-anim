import os
import docx
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# How tu run:
# uv run --with python-docx  scripts/src_to_doc.py src

def set_cell_border(cell, color_hex="D3D3D3", sz="4"):
    """为1x1表格单元格设置轻量化浅灰色边框，防止原生黑粗边框影响美观"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    border_format = {"sz": sz, "val": "single", "color": color_hex}
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = 'w:{}'.format(edge)
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, val in border_format.items():
            element.set(qn('w:{}'.format(key)), str(val))

def set_cell_background(cell, color_hex="F8F9FA"):
    """为代码框单元格设置淡淡的现代科技感背景底色"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def generate_code_appendix():
    doc = docx.Document()
    
    # 设置标准的 A4 页面页边距（上下1英寸，左右1英寸）
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. 添加附录大标题
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("附录：核心源代码清单")
    run_title.font.name = 'SimSun'
    run_title.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run_title.font.size = Pt(16)  # 三号字
    run_title.bold = True

    # 2. 遍历的源码目录及允许抽取的扩展名
    source_dir = "./src"
    valid_extensions = ('.h', '.cpp', '.hpp', '.vert', '.frag', '.glsl')

    if not os.path.exists(source_dir):
        print(f"错误: 未找到 {source_dir} 目录，请确保在项目根目录下运行此脚本。")
        return

    file_count = 0
    # 递归遍历目录下的所有文件并保持相对路径
    for root, dirs, files in os.walk(source_dir):
        # 按照字母表排序，保证附录生成顺序规范优雅
        files.sort()
        for file in files:
            if file.endswith(valid_extensions):
                full_path = os.path.join(root, file)
                # 转换为标准 Linux 相对路径显示
                relative_path = os.path.relpath(full_path, ".").replace("\\", "/")
                
                print(f"正在抽取并排版: {relative_path} ...")
                file_count += 1

                # 读取代码内容
                try:
                    with open(full_path, 'r', encoding='utf-8') as f:
                        code_content = f.read()
                except Exception as e:
                    # 容错处理：若有其他编码，尝试用 gbk
                    with open(full_path, 'r', encoding='gbk', errors='ignore') as f:
                        code_content = f.read()

                # --- 写入 Word ---
                # A. 写入相对路径标题（要求：宋体小四）
                p_path = doc.add_paragraph()
                p_path.paragraph_format.space_before = Pt(14)
                p_path.paragraph_format.space_after = Pt(4)
                p_path.paragraph_format.keep_with_next = True # 防止路径和代码框断在两个不同的页面
                
                run_path = p_path.add_run(f"文件路径: {relative_path}")
                run_path.font.name = 'SimSun'
                run_path.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                run_path.font.size = Pt(12)  # 小四
                run_path.bold = True

                # B. 创建 1*1 表格容器放置代码
                table = doc.add_table(rows=1, cols=1)
                table.autofit = False
                cell = table.cell(0, 0)
                cell.width = Inches(6.5)  # 填满正文版心
                
                set_cell_border(cell, color_hex="D3D3D3")  # 浅灰优雅边框
                set_cell_background(cell, "F8F9FA")       # 代码浅色底噪
                
                # C. 写入表格内代码（要求：宋体五号，英文使用等宽 Consolas 保证对齐）
                p_code = cell.paragraphs[0]
                p_code.paragraph_format.space_before = Pt(6)
                p_code.paragraph_format.space_after = Pt(6)
                p_code.paragraph_format.line_spacing = 1.15  # 紧凑单倍或微调行距
                
                run_code = p_code.add_run(code_content)
                run_code.font.name = 'Consolas'               # 西文使用经典等宽 Consolas 
                run_code.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun') # 中文及注释使用宋体
                run_code.font.size = Pt(10.5)                 # 五号字

    output_filename = "附录代码.docx"
    doc.save(output_filename)
    print(f"\n成功! 已成功抽取并排版完毕 {file_count} 个源文件。")
    print(f"生成的文档已保存至: {output_filename}")

if __name__ == "__main__":
    generate_code_appendix()